import shutil
import subprocess
import tempfile
from abc import ABC, abstractmethod
from io import BytesIO
from pathlib import Path

import fitz
import pytesseract
from docx import Document
from PIL import Image, UnidentifiedImageError

from app.core.errors import (
    CorruptedDocumentError,
    EmptyExtractedTextError,
    InvalidFileError,
    OCRFailureError,
)
from app.infrastructure.document.models import DocumentType
from app.infrastructure.document.normalization import normalize_text


class FileExtractor(ABC):
    document_type: DocumentType

    @abstractmethod
    def extract(self, content: bytes) -> str:
        """Extract normalized text from document bytes."""

    def _ensure_text(self, text: str) -> str:
        normalized = normalize_text(text)
        if not normalized:
            raise EmptyExtractedTextError("No readable resume text could be extracted.")
        return normalized


class TextExtractor(FileExtractor):
    document_type = DocumentType.TXT

    def extract(self, content: bytes) -> str:
        for encoding in ("utf-8-sig", "utf-8", "utf-16", "cp1252"):
            try:
                return self._ensure_text(content.decode(encoding))
            except UnicodeDecodeError:
                continue
        raise InvalidFileError("Text file encoding is not supported.")


class DOCXExtractor(FileExtractor):
    document_type = DocumentType.DOCX

    def extract(self, content: bytes) -> str:
        try:
            document = Document(BytesIO(content))
        except Exception as exc:
            raise CorruptedDocumentError("DOCX file could not be opened.") from exc

        parts: list[str] = []
        parts.extend(paragraph.text for paragraph in document.paragraphs)
        for table in document.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return self._ensure_text("\n".join(parts))


class ImageOCRExtractor(FileExtractor):
    document_type = DocumentType.PNG

    def extract(self, content: bytes) -> str:
        try:
            image = Image.open(BytesIO(content))
            image.load()
        except UnidentifiedImageError as exc:
            raise CorruptedDocumentError("Image file could not be opened.") from exc
        except Exception as exc:
            raise CorruptedDocumentError("Image file is invalid or corrupted.") from exc

        try:
            text = pytesseract.image_to_string(image)
        except pytesseract.TesseractNotFoundError as exc:
            raise OCRFailureError("Tesseract OCR is not installed or not available on PATH.") from exc
        except Exception as exc:
            raise OCRFailureError("OCR failed for the uploaded image.") from exc
        return self._ensure_text(text)


class PDFExtractor(FileExtractor):
    document_type = DocumentType.PDF

    def __init__(self, ocr_extractor: ImageOCRExtractor | None = None, min_text_chars: int = 80):
        self.ocr_extractor = ocr_extractor or ImageOCRExtractor()
        self.min_text_chars = min_text_chars

    def extract(self, content: bytes) -> str:
        try:
            document = fitz.open(stream=content, filetype="pdf")
        except Exception as exc:
            raise CorruptedDocumentError("PDF file could not be opened.") from exc

        try:
            text = "\n".join(page.get_text("text") for page in document)
            normalized = normalize_text(text)
            if len(normalized) >= self.min_text_chars:
                return normalized
            return self._extract_with_ocr(document)
        finally:
            document.close()

    def _extract_with_ocr(self, document: fitz.Document) -> str:
        parts: list[str] = []
        for page in document:
            try:
                pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                image_bytes = pixmap.tobytes("png")
                parts.append(self.ocr_extractor.extract(image_bytes))
            except EmptyExtractedTextError:
                continue
            except Exception as exc:
                raise OCRFailureError("OCR failed for scanned PDF.") from exc
        return self._ensure_text("\n".join(parts))


class DOCExtractor(FileExtractor):
    document_type = DocumentType.DOC

    def __init__(self, libreoffice_path: str | None = None):
        self.libreoffice_path = libreoffice_path

    def extract(self, content: bytes) -> str:
        soffice = self.libreoffice_path or shutil.which("soffice") or shutil.which("soffice.exe")
        if not soffice:
            raise InvalidFileError(
                "Legacy DOC extraction requires LibreOffice. Install LibreOffice or set LIBREOFFICE_PATH."
            )

        with tempfile.TemporaryDirectory(prefix="careeros-doc-") as tmp_dir:
            tmp_path = Path(tmp_dir)
            source = tmp_path / "resume.doc"
            source.write_bytes(content)
            command = [
                soffice,
                "--headless",
                "--convert-to",
                "txt:Text",
                "--outdir",
                str(tmp_path),
                str(source),
            ]
            try:
                result = subprocess.run(command, capture_output=True, text=True, timeout=45, check=False)
            except subprocess.TimeoutExpired as exc:
                raise InvalidFileError("LibreOffice DOC extraction timed out.") from exc
            except OSError as exc:
                raise InvalidFileError("LibreOffice could not be started for DOC extraction.") from exc

            output = tmp_path / "resume.txt"
            if result.returncode != 0 or not output.exists():
                raise CorruptedDocumentError("DOC file could not be converted to text.")
            return self._ensure_text(output.read_text(encoding="utf-8", errors="ignore"))
